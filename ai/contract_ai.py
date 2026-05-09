import os
import re
import json
import cv2
import fitz
import numpy as np
import pandas as pd
import pytesseract
import chromadb

from docx import Document
from pdf2image import convert_from_path
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\poppler-25.12.0\Library\bin"


# ---------- RAG / KNOWLEDGE BASE ----------
try:
    chroma_client = chromadb.PersistentClient(path="chroma_db")
    sbc_collection = chroma_client.get_or_create_collection(name="sbc")
except Exception:
    sbc_collection = None


FIDIC_KNOWLEDGE = """
Use FIDIC-inspired construction contract principles when analyzing clauses.

Important clauses to check:
- Scope of work
- Contract price and payment terms
- Payment schedule / milestones
- Time for completion / project duration
- Delay damages / penalties
- Materials and specifications
- Contractor responsibilities
- Employer responsibilities
- Variations / change orders
- Defects liability / warranty
- Termination
- Liability and indemnity
- Insurance
- Performance security
- Force majeure
- Governing law
- Dispute resolution
- Compliance with Saudi Building Code, permits, municipality approvals, and Saudi regulations
"""


# ---------- CLEAN ----------
def clean_text(text):
    text = text or ""
    return re.sub(r"\s+", " ", text).strip()


# ---------- REGEX FALLBACKS ----------
def extract_contract_value_regex(text):
    text = clean_text(text)

    patterns = [
        r"(?:contract value|project value|total value|contract amount|contract price|total price|price|amount).*?(\d[\d,\.]*)\s*(?:SAR|SR|riyal|riyals)",
        r"(?:SAR|SR)\s*(\d[\d,\.]*)",
        r"(\d[\d,\.]*)\s*(?:SAR|SR|riyal|riyals)",

        r"(?:قيمة العقد|قيمة المشروع|المبلغ الإجمالي|اجمالي قيمة العقد|إجمالي قيمة العقد|السعر الإجمالي|اجمالي المبلغ|إجمالي المبلغ|المبلغ|القيمة).*?(\d[\d,\.]*)\s*(?:ريال|ر\.س|SAR|SR)",
        r"(\d[\d,\.]*)\s*(?:ريال|ر\.س)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            return f"{value} SAR"

    return "missing"


def extract_duration_regex(text):
    text = clean_text(text)

    patterns = [
        r"(?:duration|period|completion period|project duration|time for completion).*?(\d+)\s*(months?|years?|weeks?|days?)",
        r"(\d+)\s*(months?|years?|weeks?|days?)",

        r"(?:مدة العقد|مدة المشروع|مدة التنفيذ|فترة التنفيذ|مدة الإنجاز|مدة الانجاز).*?(\d+)\s*(شهر|أشهر|اشهر|سنة|سنوات|أسبوع|اسبوع|أسابيع|اسابيع|يوم|أيام|ايام)",
        r"(\d+)\s*(شهر|أشهر|اشهر|سنة|سنوات|أسبوع|اسبوع|أسابيع|اسابيع|يوم|أيام|ايام)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return f"{match.group(1)} {match.group(2)}"

    return "missing"


def text_has_any(text, keywords):
    text = text.lower()
    return any(keyword.lower() in text for keyword in keywords)


def ensure_clause(result, clause_name, evidence_text, risk_level="Medium", status="partial", issue="", recommendation=""):
    if "clauses" not in result or not isinstance(result["clauses"], dict):
        result["clauses"] = {}

    current = result["clauses"].get(clause_name, {})
    current_status = str(current.get("status", "missing")).lower()

    if current_status in ["missing", "", "not found", "not_found", "not specified", "not_specified"]:
        result["clauses"][clause_name] = {
            "status": status,
            "risk_level": risk_level,
            "issue": issue,
            "recommendation": recommendation,
            "evidence": [evidence_text],
        }
    else:
        evidence = current.get("evidence", [])
        if not isinstance(evidence, list):
            evidence = [str(evidence)]

        if evidence_text and evidence_text not in evidence:
            evidence.append(evidence_text)

        current["evidence"] = evidence
        current["risk_level"] = current.get("risk_level") or risk_level
        current["issue"] = current.get("issue") or issue
        current["recommendation"] = current.get("recommendation") or recommendation

        result["clauses"][clause_name] = current

    return result


def infer_clauses_from_text(result, contract_text):
    text = clean_text(contract_text)
    lower = text.lower()

    if not text:
        return result

    if text_has_any(lower, ["scope of work", "construction", "structural", "excavation", "concrete", "masonry", "أعمال", "تنفيذ", "خرسانة", "بناء", "إنشاء", "انشاء"]):
        result = ensure_clause(
            result,
            "scope",
            "The document describes construction works and contractor work scope.",
            risk_level="Low",
            status="found",
            issue="",
            recommendation=""
        )

    if text_has_any(lower, ["payment", "contract value", "contract price", "amount", "SAR", "SR", "ريال", "مبلغ", "قيمة", "دفعة", "دفع"]):
        result = ensure_clause(
            result,
            "payment",
            "The document includes financial or payment-related terms.",
            risk_level="Medium",
            status="partial",
            issue="Payment terms may require clearer milestones or payment conditions.",
            recommendation="Clarify payment amount, payment method, milestones, and due dates."
        )

    if text_has_any(lower, ["milestone", "payment schedule", "installment", "دفعة", "دفعات", "جدول السداد", "مستخلص"]):
        result = ensure_clause(
            result,
            "payment_schedule",
            "The document refers to payment stages or installment-related information.",
            risk_level="Medium",
            status="partial",
            issue="Payment schedule may need clearer milestone conditions.",
            recommendation="Add a milestone-based payment schedule with approval requirements."
        )

    if text_has_any(lower, ["duration", "completion", "months", "days", "weeks", "مدة", "إنجاز", "انجاز", "شهر", "أشهر", "اشهر", "يوم", "أيام"]):
        result = ensure_clause(
            result,
            "timeline",
            "The document includes a project duration or completion period.",
            risk_level="Medium",
            status="partial",
            issue="Timeline exists but may need exact start date, completion date, and delay handling.",
            recommendation="Clarify project start date, completion date, duration, and consequences of delay."
        )

    if text_has_any(lower, ["material", "materials", "specification", "brand", "supplier", "مواد", "المواد", "مواصفات", "توريد", "اعتماد"]):
        result = ensure_clause(
            result,
            "materials",
            "The document refers to materials or technical specifications.",
            risk_level="Medium",
            status="partial",
            issue="Material specifications may not be fully detailed.",
            recommendation="Specify material type, quality, grade, supplier, and approval process."
        )

    if text_has_any(lower, ["penalty", "penalties", "delay damages", "delay", "غرامة", "غرامات", "تأخير", "تاخير", "جزاء"]):
        result = ensure_clause(
            result,
            "penalties",
            "The document refers to delay or penalty-related obligations.",
            risk_level="Medium",
            status="partial",
            issue="Delay penalties may need clearer amount, calculation method, and enforcement procedure.",
            recommendation="Clarify delay penalty value, calculation basis, grace period, and enforcement process."
        )

    if text_has_any(lower, ["warranty", "guarantee", "defects liability", "ضمان", "كفالة", "عيوب"]):
        result = ensure_clause(
            result,
            "warranty",
            "The document refers to warranty, guarantee, or defect-related responsibility.",
            risk_level="Medium",
            status="partial",
            issue="Warranty terms may need clearer duration and coverage.",
            recommendation="Specify warranty duration, covered defects, exclusions, and repair process."
        )

    if text_has_any(lower, ["termination", "terminate", "فسخ", "إنهاء", "انهاء"]):
        result = ensure_clause(
            result,
            "termination",
            "The document refers to termination or contract ending conditions.",
            risk_level="Medium",
            status="partial",
            issue="Termination procedure may require clearer notice period and consequences.",
            recommendation="Clarify termination rights, notice period, compensation, and handover obligations."
        )

    if text_has_any(lower, ["liability", "indemnity", "damage", "damages", "مسؤولية", "تعويض", "أضرار", "اضرار"]):
        result = ensure_clause(
            result,
            "liability",
            "The document refers to responsibility, damages, or compensation.",
            risk_level="Medium",
            status="partial",
            issue="Liability scope may need clearer limits and responsibilities.",
            recommendation="Clarify liability limits, indemnity obligations, and responsibility for damages."
        )

    if text_has_any(lower, ["saudi building code", "sbc", "municipality", "permit", "permits", "code", "كود البناء السعودي", "البلدية", "رخصة", "تصريح", "اشتراطات"]):
        result = ensure_clause(
            result,
            "compliance",
            "The document refers to regulatory, permit, or Saudi compliance requirements.",
            risk_level="Medium",
            status="partial",
            issue="Compliance requirements may need clearer standards and responsible party.",
            recommendation="Add clear compliance obligations for Saudi Building Code, permits, municipality approvals, and safety regulations."
        )

    if text_has_any(lower, ["dispute", "arbitration", "court", "محكمة", "تحكيم", "نزاع", "خلاف"]):
        result = ensure_clause(
            result,
            "dispute_resolution",
            "The document refers to dispute, arbitration, or court handling.",
            risk_level="Medium",
            status="partial",
            issue="Dispute resolution process may need clearer jurisdiction and steps.",
            recommendation="Specify dispute escalation steps, governing court or arbitration body, and applicable jurisdiction."
        )

    if text_has_any(lower, ["law", "governing law", "regulation", "نظام", "القانون", "الأنظمة", "الانظمة"]):
        result = ensure_clause(
            result,
            "governing_law",
            "The document refers to laws or regulations.",
            risk_level="Medium",
            status="partial",
            issue="Governing law may need explicit wording.",
            recommendation="Clearly state the governing law and applicable Saudi regulations."
        )

    return result


# ---------- RAG RETRIEVAL ----------
def retrieve_sbc_context(contract_text, top_k=4):
    if sbc_collection is None or not contract_text:
        return ""

    try:
        query_text = clean_text(contract_text)[:3000]

        embedding = client.embeddings.create(
            model="text-embedding-3-small",
            input=query_text
        ).data[0].embedding

        results = sbc_collection.query(
            query_embeddings=[embedding],
            n_results=top_k
        )

        docs = results.get("documents", [[]])[0]
        return "\n\n".join(docs)

    except Exception:
        return ""


# ---------- OCR ----------
def preprocess_image(img):
    if img is None:
        return None

    img = np.array(img)

    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    else:
        gray = img

    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh = cv2.threshold(
        blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
    )
    return thresh


def extract_text_from_image(path):
    img = cv2.imread(path)

    if img is None:
        return {"text": "", "method": "ocr_image_failed"}

    processed = preprocess_image(img)

    if processed is None:
        return {"text": "", "method": "ocr_image_failed"}

    text = pytesseract.image_to_string(processed, lang="ara+eng")
    return {"text": clean_text(text), "method": "ocr_image"}


# ---------- PDF ----------
def extract_text_from_pdf(path):
    text = ""

    try:
        doc = fitz.open(path)

        for page in doc:
            text += page.get_text() + "\n"

        doc.close()

        text = clean_text(text)

        if len(text) > 100:
            return {"text": text, "method": "pdf_text"}

    except Exception:
        pass

    try:
        images = convert_from_path(path, poppler_path=POPPLER_PATH)
        ocr_text = ""

        for img in images:
            processed = preprocess_image(img)

            if processed is not None:
                ocr_text += pytesseract.image_to_string(
                    processed,
                    lang="ara+eng"
                ) + "\n"

        ocr_text = clean_text(ocr_text)

        if ocr_text:
            return {"text": ocr_text, "method": "pdf_ocr"}

        return {"text": "", "method": "pdf_ocr_empty"}

    except Exception:
        return {"text": text, "method": "pdf_failed"}


# ---------- DOCX ----------
def extract_text_from_docx(path):
    doc = Document(path)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return {"text": clean_text(text), "method": "docx"}


# ---------- EXCEL ----------
def extract_text_from_excel(path):
    sheets = pd.read_excel(path, sheet_name=None)
    text = ""

    for sheet_name, df in sheets.items():
        text += f"\nSheet: {sheet_name}\n"
        text += df.astype(str).to_string(index=False)

    return {"text": clean_text(text), "method": "excel"}


# ---------- DATES ----------
def extract_dates_regex(text):
    patterns = [
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
        r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b",
        r"\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b",
        r"\b\d{1,2}\s+(يناير|فبراير|مارس|أبريل|ابريل|مايو|يونيو|يوليو|أغسطس|اغسطس|سبتمبر|أكتوبر|اكتوبر|نوفمبر|ديسمبر)\s+\d{4}\b",
    ]

    found = []

    for pattern in patterns:
        matches = re.findall(pattern, text, flags=re.IGNORECASE)

        for match in matches:
            if isinstance(match, tuple):
                continue
            found.append(match)

    return list(dict.fromkeys(found))


# ---------- DEFAULT ----------
def default_result():
    return {
        "contract_type": "Unknown",
        "document_type": "Unknown",
        "is_full_contract": False,
        "contract_subtype": "Unknown",
        "contract_quality_score": 0,
        "completeness_score": 0,
        "number_of_extracted_clauses": 0,

        "contract_overview": {
            "contract_type": "Unknown",
            "parties": [],
            "contract_start_date": "missing",
            "contract_duration": "missing",
            "contract_value": "missing"
        },

        "summary": [],
        "parties": [],
        "dates": [],
        "financial_terms": [],
        "construction_scope": [],
        "materials": [],
        "technical_requirements": [],
        "quality_constraints": [],
        "warranty_or_guarantees": [],
        "construction_risks": [],

        "clauses": {
            "scope": {"status": "missing", "evidence": [], "risk_level": "High", "issue": "", "recommendation": ""},
            "payment": {"status": "missing", "evidence": [], "risk_level": "High", "issue": "", "recommendation": ""},
            "payment_schedule": {"status": "missing", "evidence": [], "risk_level": "High", "issue": "", "recommendation": ""},
            "timeline": {"status": "missing", "evidence": [], "risk_level": "High", "issue": "", "recommendation": ""},
            "materials": {"status": "missing", "evidence": [], "risk_level": "High", "issue": "", "recommendation": ""},
            "warranty": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "termination": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "liability": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "confidentiality": {"status": "missing", "evidence": [], "risk_level": "Low", "issue": "", "recommendation": ""},
            "governing_law": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "dispute_resolution": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "penalties": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""},
            "renewal": {"status": "missing", "evidence": [], "risk_level": "Low", "issue": "", "recommendation": ""},
            "compliance": {"status": "missing", "evidence": [], "risk_level": "Medium", "issue": "", "recommendation": ""}
        },

        "extracted_clauses": [],
        "ambiguous_clauses": [],
        "missing_clauses": [],

        "ai_recommendations": {
            "ambiguous_clauses": [],
            "missing_core_clauses": [],
            "improvement_recommendations": []
        },

        "risks": [],
        "overall_risk": "Unknown",
        "missing_critical_sections": [],
        "recommendations": [],
        "confidence": "Low",
        "rag_used": False
    }


# ---------- PARSE ----------
def parse_json(text):
    try:
        return json.loads(text)
    except Exception:
        start = text.find("{")
        end = text.rfind("}")

        if start == -1 or end == -1:
            raise ValueError("AI did not return valid JSON.")

        return json.loads(text[start:end + 1])


# ---------- CHUNKING ----------
def split_text(text, chunk_size=8000, overlap=500):
    text = clean_text(text)

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap

        if start < 0:
            start = 0

        if start >= len(text):
            break

    return chunks


# ---------- PROMPT ----------
def build_prompt(contract_text, detected_dates, sbc_context=""):
    return f"""
You are a senior construction contract analyst and legal/technical reviewer.

Use the following FIDIC-inspired knowledge when relevant:
{FIDIC_KNOWLEDGE}

Use the following retrieved Saudi Building Code / project knowledge base context when relevant.
Do not copy it directly. Use it only to improve compliance analysis and recommendations:
{sbc_context}

The document may be Arabic or English.

IMPORTANT:
This is NOT only legal analysis.
If the document includes construction works, materials, quantities, brands, specifications, workmanship standards, or BOQ items, analyze them carefully.

STRICT RULES:
- Return ONLY valid JSON.
- Do NOT guess.
- Use ONLY the provided contract text.
- Do NOT invent missing information.
- If something is missing, write "missing".
- Evidence MUST always be a list.
- Evidence must be paraphrased summaries, NOT copied text.
- Do NOT copy long sentences from the contract.
- Summary must be specific to THIS document.
- For Arabic contracts, keep extracted Arabic names and terms as they appear.
- Extract contract_value clearly if any amount appears as the contract amount, total project value, total price, or lump sum.
- Extract contract_duration clearly if the document mentions duration, completion period, or execution period.
- Extract found and partial clauses even when they are not perfect.

Normalize party names:
- "Employer", "Owner", "Client" → always return as "Owner"
- "Contractor" stays "Contractor"

Document type rules:
- Full Construction Contract: has legal + financial + technical clauses.
- BOQ: mostly work items, quantities, prices, units, or "مقطوعية".
- Quotation: price offer, proposal, عرض سعر.
- Invoice: فاتورة, invoice, tax invoice.

Clause status rules:
- "found" = clause exists and has useful enforceable details.
- "partial" = clause exists but is vague, incomplete, unclear, or missing important details.
- "missing" = clause is not found in the contract.

PRE-DETECTED DATES:
{detected_dates}

Return ONLY valid JSON in this exact structure:

{{
  "contract_type": "",
  "document_type": "",
  "is_full_contract": false,
  "contract_subtype": "",
  "contract_value": "",
  "duration": "",
  "summary": ["", "", ""],
  "parties": [],
  "dates": [
    {{"date": "", "type": "", "evidence": []}}
  ],
  "financial_terms": [],
  "clauses": {{
    "scope": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "payment": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "payment_schedule": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "timeline": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "materials": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "warranty": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "termination": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "liability": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "confidentiality": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "governing_law": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "dispute_resolution": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "penalties": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "renewal": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}},
    "compliance": {{"status": "found or partial or missing", "risk_level": "Low or Medium or High", "issue": "", "recommendation": "", "evidence": []}}
  }},
  "construction_scope": [
    {{"category": "", "description": "", "quantity_or_area": "Not specified", "evidence": []}}
  ],
  "materials": [
    {{"material": "", "specification": "Not specified", "brand_or_supplier": "Not specified", "status": "specified or partially specified or not specified", "evidence": []}}
  ],
  "technical_requirements": [
    {{"requirement": "", "category": "", "evidence": []}}
  ],
  "quality_constraints": [
    {{"constraint": "", "reason": "", "evidence": []}}
  ],
  "warranty_or_guarantees": [
    {{"item": "", "duration": "", "evidence": []}}
  ],
  "risks": [
    {{"name": "", "level": "Low or Medium or High", "reason": "", "evidence": []}}
  ],
  "construction_risks": [
    {{"name": "", "level": "Low or Medium or High", "reason": "", "related_work": "", "evidence": []}}
  ],
  "ambiguous_clauses": [
    {{"clause": "", "issue": "", "why_it_matters": "", "recommendation": "", "evidence": []}}
  ],
  "missing_clauses": [
    {{"clause": "", "importance": "", "risk_if_missing": "", "recommendation": ""}}
  ],
  "overall_risk": "Low or Medium or High",
  "recommendations": [],
  "confidence": "Low or Medium or High"
}}

Contract text:
{contract_text}
"""


# ---------- AI SINGLE CHUNK ----------
def analyze_chunk_with_openai(chunk_text, detected_dates, sbc_context=""):
    prompt = build_prompt(chunk_text, detected_dates, sbc_context)

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a precise construction and legal contract AI. Return strict JSON only."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
    except Exception:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a precise construction and legal contract AI. Return strict JSON only."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )

    return parse_json(response.choices[0].message.content)


# ---------- MERGE ----------
def merge_status(old_status, new_status):
    rank = {
        "missing": 0,
        "not specified": 0,
        "unclear": 1,
        "partial": 2,
        "partially specified": 2,
        "found": 3,
        "specified": 3
    }

    old = (old_status or "missing").lower()
    new = (new_status or "missing").lower()

    return new_status if rank.get(new, 0) > rank.get(old, 0) else old_status


def merge_unique_list(old_list, new_list):
    old_list = old_list or []
    new_list = new_list or []

    seen = set()
    merged = []

    for item in old_list + new_list:
        key = json.dumps(item, ensure_ascii=False, sort_keys=True) if isinstance(item, dict) else str(item)

        if key not in seen:
            seen.add(key)
            merged.append(item)

    return merged


def merge_results(results):
    final = default_result()

    for result in results:
        if not result:
            continue

        for direct_key in ["contract_type", "document_type", "contract_subtype", "contract_value", "duration"]:
            if final.get(direct_key) in [None, "", "Unknown", "missing"] and result.get(direct_key):
                final[direct_key] = result.get(direct_key)

        if result.get("is_full_contract"):
            final["is_full_contract"] = True

        for key in [
            "summary",
            "parties",
            "dates",
            "financial_terms",
            "construction_scope",
            "materials",
            "technical_requirements",
            "quality_constraints",
            "warranty_or_guarantees",
            "risks",
            "construction_risks",
            "recommendations",
            "ambiguous_clauses",
            "missing_clauses"
        ]:
            final[key] = merge_unique_list(final.get(key), result.get(key))

        for clause_name, clause_data in result.get("clauses", {}).items():
            if clause_name not in final["clauses"]:
                final["clauses"][clause_name] = {
                    "status": "missing",
                    "evidence": [],
                    "risk_level": "",
                    "issue": "",
                    "recommendation": ""
                }

            old_status = final["clauses"][clause_name].get("status")
            new_status = clause_data.get("status")

            final["clauses"][clause_name]["status"] = merge_status(old_status, new_status)
            final["clauses"][clause_name]["evidence"] = merge_unique_list(
                final["clauses"][clause_name].get("evidence"),
                clause_data.get("evidence")
            )

            for field in ["risk_level", "issue", "recommendation"]:
                if clause_data.get(field):
                    final["clauses"][clause_name][field] = clause_data.get(field)

    return final


# ---------- POST PROCESS ----------
def enforce_document_type(result, text):
    t = text.lower()
    current_type = (result.get("document_type") or result.get("contract_type") or "").lower()

    if not current_type or current_type == "unknown":
        if "مقطوعية" in text or "boq" in t or "bill of quantities" in t:
            result["document_type"] = "BOQ"
            result["contract_type"] = "BOQ / Bill of Quantities"
            result["is_full_contract"] = False

        elif "عرض سعر" in text or "quotation" in t or "price offer" in t:
            result["document_type"] = "Quotation"
            result["contract_type"] = "Quotation / Price Offer"
            result["is_full_contract"] = False

        elif "فاتورة" in text or "invoice" in t:
            result["document_type"] = "Invoice"
            result["contract_type"] = "Invoice"
            result["is_full_contract"] = False

        elif "عقد" in text or "agreement" in t or "contract" in t:
            result["document_type"] = "Full Construction Contract"
            result["contract_type"] = "Construction Contract"
            result["is_full_contract"] = True

        else:
            result["document_type"] = "Unknown"
            result["contract_type"] = "Unknown"

    if "مقطوعية" in text and "عقد" not in text:
        result["document_type"] = "BOQ"
        result["contract_type"] = "BOQ / Bill of Quantities"
        result["is_full_contract"] = False

    return result


def clean_materials(result):
    materials = result.get("materials", [])
    cleaned = []

    for material in materials:
        if not isinstance(material, dict):
            continue

        name = material.get("material") or material.get("name") or "Unknown material"
        spec = material.get("specification") or "Not specified"
        brand = material.get("brand_or_supplier") or "Not specified"
        evidence = material.get("evidence") or []
        status = material.get("status") or "partially specified"

        if name != "Unknown material" and spec == "Not specified" and brand == "Not specified":
            status = "partially specified"

        cleaned.append({
            "material": name,
            "specification": spec,
            "brand_or_supplier": brand,
            "status": status,
            "evidence": evidence if isinstance(evidence, list) else [str(evidence)]
        })

    result["materials"] = cleaned
    return result


def normalize_clauses(result):
    default_clauses = default_result()["clauses"]
    clauses = result.get("clauses", {})

    if not isinstance(clauses, dict):
        clauses = {}

    for clause_name, default_clause in default_clauses.items():
        if clause_name not in clauses or not isinstance(clauses.get(clause_name), dict):
            clauses[clause_name] = default_clause.copy()

        clauses[clause_name]["status"] = clauses[clause_name].get("status") or "missing"
        clauses[clause_name]["evidence"] = clauses[clause_name].get("evidence") or []
        clauses[clause_name]["risk_level"] = clauses[clause_name].get("risk_level") or default_clause["risk_level"]
        clauses[clause_name]["issue"] = clauses[clause_name].get("issue") or ""
        clauses[clause_name]["recommendation"] = clauses[clause_name].get("recommendation") or ""

    result["clauses"] = clauses
    return result


def normalize_key(key):
    return str(key).strip().lower().replace(" ", "_").replace("-", "_")


def apply_score(result):
    raw_clauses = result.get("clauses", {}) or {}

    clauses = {
        normalize_key(k): v
        for k, v in raw_clauses.items()
        if isinstance(v, dict)
    }

    weights = {
        "scope": 10,
        "payment": 8,
        "payment_schedule": 10,
        "timeline": 10,
        "materials": 8,
        "penalties": 7,
        "warranty": 7,
        "termination": 7,
        "liability": 7,
        "governing_law": 6,
        "dispute_resolution": 7,
        "compliance": 8,
    }

    critical_clauses = [
        "scope",
        "payment",
        "payment_schedule",
        "timeline",
        "termination",
        "dispute_resolution",
        "compliance"
    ]

    score = 0
    missing = []

    for clause_name, weight in weights.items():
        key = normalize_key(clause_name)
        clause = clauses.get(key)

        if not clause:
            if clause_name in critical_clauses:
                missing.append(clause_name)
            continue

        status = str(clause.get("status", "missing")).strip().lower()
        risk = str(clause.get("risk_level", "medium")).strip().lower()

        if status in ["missing", "not found", "not_found", "not specified", "not_specified"]:
            clause_score = 0
            if clause_name in critical_clauses:
                missing.append(clause_name)

        elif risk == "low":
            clause_score = weight

        elif risk == "medium":
            clause_score = weight * 0.6

        elif risk == "high":
            clause_score = weight * 0.3

        else:
            if status in ["found", "specified"]:
                clause_score = weight * 0.8
            elif status in ["partial", "partially specified", "unclear"]:
                clause_score = weight * 0.5
            else:
                clause_score = weight * 0.4

        score += clause_score

    score = round(score)

    found_count = 0
    for clause in clauses.values():
        status = str(clause.get("status", "")).strip().lower()
        if status not in ["missing", "not found", "not_found", "not specified", "not_specified"]:
            found_count += 1

    if score == 0 and found_count > 0:
        score = min(70, found_count * 6)

    doc_type = (result.get("document_type") or "").lower()

    if "boq" in doc_type or "quotation" in doc_type or "invoice" in doc_type:
        result["is_full_contract"] = False
        if score > 70:
            score = 70

    if score >= 85:
        risk = "Low"
    elif score >= 60:
        risk = "Medium"
    else:
        risk = "High"

    result["contract_quality_score"] = score
    result["completeness_score"] = score
    result["overall_risk"] = risk
    result["missing_critical_sections"] = missing

    return result


def ensure_recommendations(result):
    recommendations = result.get("recommendations") or []
    missing = result.get("missing_critical_sections") or []

    default_recs = {
        "scope": "Add a clear and detailed scope of work.",
        "payment": "Add clear contract value and payment obligations.",
        "payment_schedule": "Add a milestone-based payment schedule.",
        "timeline": "Add project start date, completion date, and delay rules.",
        "materials": "Clarify whether materials are included and specify material quality, grade, and supplier.",
        "termination": "Add a clear termination clause.",
        "dispute_resolution": "Add a clear dispute resolution clause.",
        "compliance": "Add compliance with Saudi Building Code, permits, safety regulations, and municipality requirements."
    }

    for item in missing:
        rec = default_recs.get(item)

        if rec and rec not in recommendations:
            recommendations.append(rec)

    if not recommendations:
        recommendations.append(
            "Improve the document by adding clearer legal, financial, and technical details."
        )

    result["recommendations"] = recommendations
    return result


# ---------- DISPLAY SECTIONS ----------
def build_contract_overview(result, contract_text=""):
    overview = {
        "contract_type": result.get("contract_type", "Unknown"),
        "parties": [],
        "contract_start_date": "missing",
        "contract_duration": "missing",
        "contract_value": "missing"
    }

    for party in result.get("parties", []):
        if isinstance(party, dict):
            role = party.get("role", "Party")
            name = party.get("name", "Not specified")

            formatted = f"{role}: {name}"

            if party.get("commercial_registration"):
                formatted += f" | CR: {party.get('commercial_registration')}"

            overview["parties"].append(formatted)

        else:
            overview["parties"].append(str(party))

    for item in result.get("dates", []):
        if isinstance(item, dict):
            date_type = str(item.get("type", "")).lower()

            if any(word in date_type for word in [
                "start",
                "commencement",
                "signing",
                "signature",
                "contract signing",
                "effective",
                "توقيع",
                "بداية",
                "بدء"
            ]):
                overview["contract_start_date"] = item.get("date", "missing")
                break

    direct_value = result.get("contract_value")
    if direct_value and direct_value not in ["missing", "Unknown", ""]:
        overview["contract_value"] = direct_value

    if overview["contract_value"] == "missing":
        for item in result.get("financial_terms", []):
            if isinstance(item, dict):
                possible_keys = [
                    "total_contract_value",
                    "contract_value",
                    "total_value",
                    "total_price",
                    "amount",
                    "price"
                ]

                for key in possible_keys:
                    value = item.get(key)

                    if value and value not in ["missing", "Unknown", ""]:
                        overview["contract_value"] = value
                        break

                if overview["contract_value"] != "missing":
                    break

    if overview["contract_value"] == "missing":
        overview["contract_value"] = extract_contract_value_regex(contract_text)

    direct_duration = result.get("duration") or result.get("contract_duration") or result.get("project_duration")
    if direct_duration and direct_duration not in ["missing", "Unknown", ""]:
        overview["contract_duration"] = direct_duration

    if overview["contract_duration"] == "missing":
        for clause in result.get("clauses", {}).values():
            if isinstance(clause, dict):
                evidence = " ".join(clause.get("evidence", []))

                if any(word in evidence.lower() for word in [
                    "duration",
                    "completion",
                    "months",
                    "days",
                    "weeks",
                    "مدة",
                    "إنجاز",
                    "انجاز",
                    "شهر",
                    "أشهر",
                    "اشهر"
                ]):
                    overview["contract_duration"] = evidence
                    break

    if overview["contract_duration"] == "missing":
        overview["contract_duration"] = extract_duration_regex(contract_text)

    result["contract_overview"] = overview
    result["contract_value"] = overview["contract_value"]
    result["duration"] = overview["contract_duration"]

    return result


def build_extracted_clauses_table(result):
    rows = []

    for clause_name, clause_data in result.get("clauses", {}).items():
        status = str(clause_data.get("status", "missing")).lower()

        if status in ["missing", "not found", "not_found", "not specified", "not_specified"]:
            continue

        evidence = clause_data.get("evidence", [])
        if not isinstance(evidence, list):
            evidence = [str(evidence)]

        extracted_text = (
            evidence[0]
            if evidence
            else clause_data.get("issue")
            or clause_data.get("recommendation")
            or "Clause detected."
        )

        rows.append({
            "clause_category": clause_name.replace("_", " ").title(),
            "status": clause_data.get("status", "missing"),
            "risk_level": clause_data.get("risk_level", "Unknown"),
            "extracted_text": extracted_text,
            "issue": clause_data.get("issue", ""),
            "recommendation": clause_data.get("recommendation", "")
        })

    result["extracted_clauses"] = rows
    result["number_of_extracted_clauses"] = len(rows)

    return result


def build_missing_and_recommendations(result):
    clauses = result.get("clauses", {})

    required_clauses = {
        "scope": "Defines the work included in the project.",
        "payment": "Defines the contract value and payment obligation.",
        "payment_schedule": "Connects payments to clear milestones.",
        "timeline": "Defines duration, start date, completion date, and delay handling.",
        "materials": "Defines material quality, specifications, and approvals.",
        "warranty": "Defines responsibility for defects after handover.",
        "termination": "Defines how either party may end the contract.",
        "liability": "Defines responsibility for damages, negligence, and losses.",
        "governing_law": "Defines the law that controls the contract.",
        "dispute_resolution": "Defines how disputes will be resolved.",
        "penalties": "Defines consequences of delay or non-performance.",
        "compliance": "Ensures compliance with Saudi Building Code, permits, and regulations.",
        "confidentiality": "Protects confidential project and commercial information."
    }

    ambiguous = []
    missing = []
    improvement_recommendations = []

    for clause_name, importance in required_clauses.items():
        data = clauses.get(clause_name, {})
        status = str(data.get("status", "missing")).lower()
        risk = str(data.get("risk_level", "")).lower()

        readable = clause_name.replace("_", " ").title()

        if status in ["missing", "not found", "not_found", "not specified", "not_specified"]:
            rec = f"Add a clear {readable} clause with specific responsibilities, requirements, and consequences."

            missing.append({
                "clause": readable,
                "importance": importance,
                "risk_if_missing": "May create contractual uncertainty or disputes.",
                "recommendation": rec
            })

            improvement_recommendations.append(rec)

        elif risk in ["medium", "high"]:
            rec = data.get("recommendation") or f"Clarify {readable} with measurable terms, responsible parties, dates, and consequences."

            ambiguous.append({
                "clause": readable,
                "issue": data.get("issue") or f"{readable} exists but may create risk due to unclear or incomplete terms.",
                "why_it_matters": importance,
                "recommendation": rec
            })

            improvement_recommendations.append(rec)

    result["ai_recommendations"] = {
        "ambiguous_clauses": ambiguous,
        "missing_core_clauses": missing,
        "improvement_recommendations": improvement_recommendations
    }

    result["missing_clauses"] = missing
    result["ambiguous_clauses"] = ambiguous

    return result


def finalize_display_sections(result, contract_text="", rag_used=False):
    result["completeness_score"] = result.get("contract_quality_score", 0)
    result["rag_used"] = rag_used

    result = build_contract_overview(result, contract_text)
    result = build_extracted_clauses_table(result)
    result = build_missing_and_recommendations(result)

    return result


# ---------- AI MAIN ----------
def analyze_with_openai(contract_text):
    contract_text = clean_text(contract_text)
    detected_dates = extract_dates_regex(contract_text)

    if not contract_text:
        result = default_result()
        result["recommendations"] = ["No readable text was extracted from the file."]
        result = finalize_display_sections(result, contract_text, rag_used=False)
        return result

    sbc_context = retrieve_sbc_context(contract_text)
    rag_used = bool(sbc_context)

    chunks = split_text(contract_text, chunk_size=8000, overlap=500)
    results = []

    for chunk in chunks:
        try:
            results.append(analyze_chunk_with_openai(chunk, detected_dates, sbc_context))
        except Exception:
            continue

    if not results:
        result = default_result()
        result["recommendations"] = ["AI analysis failed. Try uploading a clearer file."]
        result = infer_clauses_from_text(result, contract_text)
        result = apply_score(result)
        result = ensure_recommendations(result)
        result = finalize_display_sections(result, contract_text, rag_used=rag_used)
        return result

    if len(results) == 1:
        result = default_result()
        result.update(results[0])
    else:
        result = merge_results(results)

    result = enforce_document_type(result, contract_text)
    result = normalize_clauses(result)
    result = infer_clauses_from_text(result, contract_text)
    result = clean_materials(result)
    result = apply_score(result)
    result = ensure_recommendations(result)
    result = finalize_display_sections(result, contract_text, rag_used=rag_used)

    return result


# ---------- MAIN ----------
def analyze_contract(file_path=None, text=None):
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            extracted = extract_text_from_pdf(file_path)

        elif ext == ".docx":
            extracted = extract_text_from_docx(file_path)

        elif ext == ".xlsx":
            extracted = extract_text_from_excel(file_path)

        elif ext in [".png", ".jpg", ".jpeg"]:
            extracted = extract_text_from_image(file_path)

        else:
            extracted = {"text": "", "method": "unsupported_file"}

        analysis = analyze_with_openai(extracted["text"])
        analysis["extraction_method"] = extracted["method"]
        return analysis

    if text:
        analysis = analyze_with_openai(text)
        analysis["extraction_method"] = "direct_text"
        return analysis

    result = default_result()
    result["recommendations"] = ["No contract provided."]
    result["extraction_method"] = "none"
    result = finalize_display_sections(result, "", rag_used=False)
    return result